import os
# -*- coding: UTF8 -*-
from docx import Document
from docx.shared import Pt
import requests
from bs4 import BeautifulSoup
import numpy as np

doc = Document()

path = "/Users/lennawy/Desktop/test/docxt/"

num = 0
check = np.zeros(200)
for page in range(12,14):
    response0 = requests.get(
        "https://www.daad.org.tw/de/ueber-uns/aktuelles/page/"+str(page)+"/")
    soup0 = BeautifulSoup(response0.text, "html.parser")
    teaser_results = soup0.find_all("div", class_="teaser two-column-teaser")
    for tesear_result in teaser_results:
        num+=1
        st0 = tesear_result.select_one("a")
        title=st0.get("title")
        st=st0.get("href")
        print(st)
        response = requests.get(
        st)
        soup = BeautifulSoup(response.text, "html.parser")
        result = soup.find("div", class_="content-area row").find("div").find("div")
        img_link0 = result.find("p").select_one("img")
        if img_link0:
            check[num]=1
            img_link=img_link0.get("src")
            print(img_link)
            if not os.path.exists("images"):
                os.mkdir("images")  # 建立資料夾
                print("directory built!")
            img = requests.get(img_link)  # download image
            pic_out = open(path+"images/"  + str(num) + ".jpg", "wb")  # 開啟資料夾及命名圖片檔
            pic_out.write(img.content)  # 寫入圖片的二進位碼
        
        paras = result.select("p")
        doc.add_heading(title)
        if check[num] == 1:
            doc.add_picture(path + 'images/'+str(num)+'.jpg', width=Pt(500))
        for para in paras:
            doc.add_paragraph(para.getText())  # 輸出排版後的HTML內容
        doc.add_paragraph('\n')
        info = result.select_one("div", class_="infobox infobox-full")#.select_one("p").select_one("a").get("href")
        doc.add_paragraph(str(info))
        doc.add_paragraph('\n')
        doc.add_paragraph('\n')
        doc_table = doc.add_page_break()

print(num)

doc.save(path + "new.doc")
